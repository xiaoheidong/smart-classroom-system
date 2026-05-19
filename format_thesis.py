"""
论文排版脚本 - 将 Markdown 初稿转换为符合格式要求的 .docx 文件
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import markdown

# ============ 字号映射 ============
FONT_SIZES = {
    '初号': 42, '小初': 36,
    '一号': 26, '小一': 24,
    '二号': 22, '小二': 18,
    '三号': 16, '小三': 15,
    '四号': 14, '小四': 12,
    '五号': 10.5, '小五': 9,
}

# ============ 颜色 ============
BLACK = RGBColor(0, 0, 0)

def set_run_font(run, cn_font='宋体', en_font='Times New Roman'):
    """设置中英文字体"""
    run.font.name = en_font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cn_font)

def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule='exact'):
    """设置段落间距"""
    pPr = paragraph.paragraph_format
    if before is not None:
        pPr.space_before = Pt(before)
    if after is not None:
        pPr.space_after = Pt(after)
    if line is not None:
        if line_rule == 'single':
            pPr.line_spacing = Pt(line)
        elif line_rule == 'exact':
            pPr.line_spacing = Pt(line)
            pPr.line_spacing_rule = 1  # exact

def set_paragraph_alignment(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = alignment

def add_formatted_paragraph(doc, text, font_cn='宋体', font_en='Times New Roman', 
                           size='小四', bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before=0, after=0, line=None, line_rule='exact',
                           page_break_before=False):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.clear()
    
    # 分页控制
    if page_break_before:
        p.paragraph_format.page_break_before = True
    
    # 对齐
    p.alignment = alignment
    
    # 间距
    set_paragraph_spacing(p, before=before, after=after, line=line, line_rule=line_rule)
    
    # 文本
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en)
    run.font.size = Pt(FONT_SIZES[size])
    run.bold = bold
    
    return p

# ============ 创建文档 ============
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
r = style.element.rPr
if r is None:
    r = parse_xml(f'<w:rPr {nsdecls("w")}/>')
r.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(FONT_SIZES['小四'])
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(20)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

print("开始排版...")

# ==================== 1. 封面 ====================
print("处理封面...")

cover_items = [
    ("本科学生毕业论文（设计）", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("", '楷体', 'Times New Roman', '小三', False, WD_ALIGN_PARAGRAPH.CENTER),
    ("题目（中文）：基于深度学习的智慧教室课堂行为分析与考勤系统设计与实现", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("", '楷体', 'Times New Roman', '小三', False, WD_ALIGN_PARAGRAPH.CENTER),
    ("（英文）：Design and Implementation of a Smart Classroom System for Classroom Behavior Analysis and Attendance Based on Deep Learning", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("", '楷体', 'Times New Roman', '小三', False, WD_ALIGN_PARAGRAPH.CENTER),
    ("姓　　名：【请填写】", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("学　　号：【请填写】", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("院（系）：信息工程学院", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("专业、年级：数据科学与大数据技术 20【请填写】级", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("指导教师：【请填写】", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
    ("", '楷体', 'Times New Roman', '小三', False, WD_ALIGN_PARAGRAPH.CENTER),
    ("日　　期：【请填写】年【请填写】月【请填写】日", '楷体', 'Times New Roman', '小三', True, WD_ALIGN_PARAGRAPH.CENTER),
]

for text, cn_font, en_font, size, bold, align in cover_items:
    p = doc.add_paragraph()
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(30)  # 单倍行距约30pt
    
    if text:
        run = p.add_run(text)
        set_run_font(run, cn_font, en_font)
        run.font.size = Pt(FONT_SIZES[size])
        run.bold = bold

# 封面后分页
doc.add_page_break()

# ==================== 2. 诚信声明 ====================
print("处理诚信声明...")

p = add_formatted_paragraph(doc, "湖南科技学院本科毕业论文（设计）诚信声明",
                           '黑体', 'Times New Roman', '三号', True, WD_ALIGN_PARAGRAPH.CENTER,
                           before=0, after=0, line=20)
doc.add_paragraph()  # 空行

statement = "本人郑重声明：所呈交的本科毕业论文（设计），是本人在指导老师的指导下，独立进行研究工作所取得的成果，成果不存在知识产权争议，除文中已经注明引用的内容外，本论文不含任何其他个人或集体已经发表或撰写过的作品成果。对本文的研究做出重要贡献的个人和集体均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。"
add_formatted_paragraph(doc, statement, '宋体', 'Times New Roman', '小四', False,
                       WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)
doc.add_paragraph()  # 空行

add_formatted_paragraph(doc, "本科毕业论文（设计）作者签名：________________",
                       '宋体', 'Times New Roman', '小四', False, WD_ALIGN_PARAGRAPH.LEFT,
                       before=0, after=0, line=20)
add_formatted_paragraph(doc, "二〇【请填写】年 【请填写】月 【请填写】日",
                       '宋体', 'Times New Roman', '小四', False, WD_ALIGN_PARAGRAPH.LEFT,
                       before=0, after=0, line=20)

doc.add_page_break()

# ==================== 3. 目录 ====================
print("处理目录...")

p = add_formatted_paragraph(doc, "目 录", '黑体', 'Times New Roman', '三号', False,
                           WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=25)

toc_items = [
    "1 绪论 ........................................................................ i",
    "    1.1 研究背景与意义 ................................................... i",
    "    1.2 国内外研究现状 ................................................... iii",
    "    1.3 主要研究内容与方法 ............................................ vii",
    "    1.4 论文结构安排 ...................................................... viii",
    "2 相关技术与理论基础 .................................................... ix",
    "    2.1 Python 与 OpenCV 视频处理 ................................. ix",
    "    2.2 PyTorch 与卷积神经网络 ..................................... x",
    "    2.3 目标检测与 YOLO 系列 ........................................ xi",
    "    2.4 行为识别与残差网络 ........................................... xiii",
    "    2.5 人脸识别与特征对齐 ........................................... xiv",
    "    2.6 图形界面与多线程 ................................................ xv",
    "    2.7 嵌入式数据库 ...................................................... xvi",
    "    2.8 大模型 API 与文本上下文（拓展） ....................... xvi",
    "    2.9 本章小结 ............................................................. xvii",
    "3 系统需求分析 ............................................................. xvii",
    "    3.1 可行性分析 .......................................................... xvii",
    "    3.2 业务流程与角色 ................................................... xix",
    "    3.3 功能需求 .............................................................. xx",
    "    3.4 非功能需求 .......................................................... xxiii",
    "    3.5 本章小结 ............................................................. xxiv",
    "4 系统总体设计 ............................................................. xxv",
    "    4.1 设计原则 .............................................................. xxv",
    "    4.2 逻辑架构 .............................................................. xxvi",
    "    4.3 功能模块 ............................................................. xxvii",
    "    4.4 关键业务流程 ...................................................... xxix",
    "    4.5 数据库设计 .......................................................... xxxi",
    "    4.6 本章小结 ............................................................. xxxiii",
    "5 系统实现 ................................................................... xxxiv",
    "    5.1 工程结构与配置 .................................................. xxxiv",
    "    5.2 人体检测与跟踪 .................................................. xxxv",
    "    5.3 行为分类推理 ..................................................... xxxvi",
    "    5.4 异常判定与可视化 ............................................. xxxvii",
    "    5.5 异常行为抓拍存证 ............................................. xxxviii",
    "    5.6 图形界面与视频源 .............................................. xxxix",
    "    5.7 人脸与考勤 .......................................................... xl",
    "    5.8 智能问答与 DeepSeek 接口 ................................. xli",
    "    5.9 本章小结 ............................................................. xlii",
    "6 系统测试与实验 ......................................................... xliii",
    "    6.1 测试环境与方法 .................................................. xliii",
    "    6.2 功能测试 .............................................................. xliv",
    "    6.3 行为分类训练实验 ............................................... xlv",
    "    6.4 结果分析 ............................................................. xlviii",
    "    6.5 本章小结 ............................................................. xlix",
    "7 总结与展望 ................................................................ l",
    "    7.1 总结 ..................................................................... l",
    "    7.2 展望 ..................................................................... li",
    "参考文献 ...................................................................... lii",
    "致 谢 ......................................................................... liv",
]

for item in toc_items:
    p = add_formatted_paragraph(doc, item, '宋体', 'Times New Roman', '小四', False,
                               WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=25)

doc.add_page_break()

# ==================== 4. 摘要 ====================
print("处理中文摘要...")

p = add_formatted_paragraph(doc, "摘 要", '黑体', 'Times New Roman', '四号', False,
                           WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=20)
doc.add_paragraph()  # 空行

# 摘要正文
abstract_paras = [
    "随着教育数字化进程加快，借助视频开展课堂观察与学习行为分析逐渐成为研究热点。纯人工巡课依赖经验判断，往往难以兼顾全班、难以留下可复核的量化痕迹。引入计算机视觉可在减轻教师文书负担的同时补充客观线索，但须同时权衡识别精度、算力开销与隐私合规。笔者在本地单机前提下研制了一套深度学习驱动的智慧教室课堂行为分析与考勤系统：输入可为 USB 摄像头实时流或磁盘上的教学录像，流水线串联人体检测、基于裁剪图的行为分类以及人脸识别，实现多类课堂姿态的在线判别、异常提示与刷脸考勤。",
    "客户端以 PyQt5 搭建，主窗口划分为四个分页：作弊检测（课堂行为分析）、动态点名、人脸注册与智能问答。检测环节调用 Ultralytics YOLOv8n 获取行人框，裁剪后送入以 ResNet18 为骨干的分类器，得到八类行为语义；依据配置文件中划分的正常/异常集合及置信度门限，规则层对睡觉、持机使用等行为累计并触发告警。针对常见笔记本独显容量，辅以跳帧与训练时较小批量等手段约束显存；推理设备在无 CUDA 或 CPU 版 PyTorch 时回退至 CPU。人脸分支在 dlib 可用时写入 128 维描述子；若编译或依赖受限，则切换 facenet-pytorch（MTCNN 对齐 + InceptionResnetV1 嵌入），以减轻环境差异带来的复现障碍。",
    "对睡觉、使用手机等异常，系统在满足持续时长或连续帧数（可配置）等条件时自动抓拍画面并写入 SQLite 的 behavior_evidence 表，图片落盘于项目 captures/behavior/，实现可复核的行为存证；「使用手机」类事件另设冷却时间，避免同一跟踪目标短时间重复存证。智能问答分页在配置环境变量 DEEPSEEK_API_KEY 后，可将当日考勤摘要与抓拍元数据（纯文本）经 HTTPS 提交至 DeepSeek 兼容接口，由大模型生成课堂概况类回答；不向公网传输原始课堂图像，以降低隐私风险。",
    "动态点名采用场次（session）语义：每次「开始点名」创建一条考勤场次记录，「停止」时结束该场并固化本场签到；未到名单为「全班已注册学生减去本场已签到」。记录写入 SQLite（学生表、考勤表及场次表等）。可选地，在配置飞书自定义机器人 Webhook 与加签密钥后，本场结束时将本场起止时间、应到/实到/未到及未到名单推送至指定群聊，便于教学管理留痕。主窗口在切换 Tab 时停止其它分页的视频与推理线程，避免多路同时打开同一摄像头导致进程异常退出。",
    "行为分类所用数据为自建课堂图像集，已划分训练、验证与测试子集；训练阶段使用类别加权交叉熵抑制长尾、AdamW 配合余弦退火，本文归档的一次实验为 20 epoch、batch=8。依据 training_history.json，验证集最高分类准确率约 93.71%，训练损失走低过程中训练集准确率同步抬升。系统经联调后支持摄像头演示与离线文件回放。文中另单列一节讨论教室场景视频采集的伦理边界与数据安全注意点。",
    "后文各章依次交代理论预备、需求与总体方案、模块实现、基于日志的实验与不足；英文摘要与关键词与之对应，便于检索。需要说明的是，摘要中的百分比与损失数值均来自指定实验目录下的 JSON 与元数据文件，若读者更换数据划分或随机种子，数值会随之变化，复现时请以本地日志为准。",
]

for para_text in abstract_paras:
    add_formatted_paragraph(doc, para_text, '宋体', 'Times New Roman', '小四', False,
                           WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)

# 关键词
p = add_formatted_paragraph(doc, "【关键词】", '黑体', 'Times New Roman', '小四', True,
                           WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)
# 添加关键词内容（同一行）
run = p.add_run("智慧教室；深度学习；YOLOv8；行为识别；人脸识别；考勤系统；PyQt5")
set_run_font(run, '宋体', 'Times New Roman')
run.font.size = Pt(FONT_SIZES['小四'])
run.bold = False

doc.add_page_break()

# ==================== 5. 英文摘要 ====================
print("处理英文摘要...")

p = add_formatted_paragraph(doc, "Abstract", '黑体', 'Times New Roman', '四号', False,
                           WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=20)
doc.add_paragraph()  # 空行

abstract_en_paras = [
    "Classroom video is increasingly used to support learning analytics under ongoing digitalization of education. This work develops a locally deployed desktop system that chains person detection, crop-level action classification, and face recognition for behavior cues, rule-based alerts, and attendance on commodity PCs.",
    "The GUI is built with PyQt5 with four tabs: behavior analysis, roll-call, face enrollment, and LLM Q&A. YOLOv8n proposes person boxes; ResNet18 classifies crops into eight action labels. A lightweight rule layer thresholds softmax scores against configured normal/abnormal sets to flag sleeping and phone use among others. Frame skipping limits per-second compute. Behavior evidence (snapshots and metadata) is stored in SQLite (behavior_evidence) under configurable duration/consecutive-frame and cooldown rules; images are saved under captures/behavior/. The face pipeline prefers dlib 128-D embeddings and falls back to facenet-pytorch when needed. SQLite also persists encodings, session-scoped attendance rows, and session metadata. Optional Feishu (Lark) bot webhooks push end-of-session summaries (times and absentees). DeepSeek-compatible chat APIs (via DEEPSEEK_API_KEY) support text-only daily summaries without uploading raw frames. Tab switching stops other tabs' camera threads to avoid device contention.",
    "A self-collected dataset is split into train/validation/test portions. One logged training run uses 20 epochs, batch 8, AdamW, cosine scheduling, and class-weighted cross-entropy; peak validation accuracy is about 93.71% in training_history.json. Optional ImageNet initialization is controlled by --pretrained. Ethical limits of classroom video analytics are discussed.",
]

for para_text in abstract_en_paras:
    add_formatted_paragraph(doc, para_text, 'Times New Roman', 'Times New Roman', '小四', False,
                           WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)

# Keywords
p = add_formatted_paragraph(doc, "【Keywords】", '黑体', 'Times New Roman', '小四', True,
                           WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)
run = p.add_run(" smart classroom; deep learning; YOLOv8; action recognition; face recognition; attendance system; PyQt5")
set_run_font(run, 'Times New Roman', 'Times New Roman')
run.font.size = Pt(FONT_SIZES['小四'])
run.bold = False

doc.add_page_break()

# ==================== 6. 正文 ====================
print("处理正文...")

# 读取 markdown 文件并解析
with open(r'd:\SY_BYSJ\smart_classroom2\docs\毕业论文初稿_智慧教室系统.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# 跳过封面、诚信声明、目录、摘要部分，从正文开始
body_start = md_content.find('# 1 绪论')
body_content = md_content[body_start:]

# 解析 markdown
lines = body_content.split('\n')
i = 0

while i < len(lines):
    line = lines[i]
    
    # 一级标题 # 
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = add_formatted_paragraph(doc, title, '宋体', 'Times New Roman', '三号', True,
                                   WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20,
                                   page_break_before=True)
        i += 1
        continue
    
    # 二级标题 ## 
    elif line.startswith('## ') and not line.startswith('### '):
        title = line[3:].strip()
        # 跳过不需要格式化的标题（摘要等已处理）
        if title in ['摘要', 'Abstract', '参考文献', '致 谢']:
            i += 1
            continue
        p = add_formatted_paragraph(doc, title, '黑体', 'Times New Roman', '四号', False,
                                   WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=20)
        i += 1
        continue
    
    # 三级标题 ### 
    elif line.startswith('### '):
        title = line[4:].strip()
        p = add_formatted_paragraph(doc, title, '黑体', 'Times New Roman', '小四', False,
                                   WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=20)
        i += 1
        continue
    
    # 表格
    elif line.startswith('|') and '---' in lines[i+1] if i+1 < len(lines) else False:
        # 解析 markdown 表格
        table_header = line
        header_cells = [cell.strip() for cell in table_header.split('|')[1:-1]]
        i += 2  # 跳过分隔行
        
        # 创建表格
        table = doc.add_table(rows=1, cols=len(header_cells))
        table.style = 'Table Grid'
        
        # 设置表头
        for j, cell_text in enumerate(header_cells):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, '宋体', 'Times New Roman')
            run.font.size = Pt(FONT_SIZES['五号'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 读取表格行
        while i < len(lines) and lines[i].startswith('|'):
            row_data = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                set_run_font(run, '宋体', 'Times New Roman')
                run.font.size = Pt(FONT_SIZES['五号'])
            i += 1
        
        # 添加表名（如果前面有行）
        doc.add_paragraph()  # 表格后空行
        continue
    
    # 空行
    elif line.strip() == '' or line.strip() == '---':
        i += 1
        continue
    
    # 正文段落
    else:
        # 清理 markdown 格式
        text = line.strip()
        if text:
            # 移除粗体标记 **text**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # 移除引用标记 [1]
            # 移除链接标记 [text](url)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            # 移除行内代码标记
            text = re.sub(r'`(.+?)`', r'\1', text)
            
            p = add_formatted_paragraph(doc, text, '宋体', 'Times New Roman', '小四', False,
                                       WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)
        i += 1

print("正文处理完成...")

# ==================== 7. 参考文献 ====================
print("处理参考文献...")

doc.add_page_break()

p = add_formatted_paragraph(doc, "参考文献", '黑体', 'Times New Roman', '四号', False,
                           WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=20)
doc.add_paragraph()  # 空行

references = [
    "[1] Redmon J, Divvala S, Girshick R, Farhadi A. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
    "[2] Redmon J, Farhadi A. YOLO9000: Better, Faster, Stronger[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Honolulu: IEEE, 2017: 6517-6525.",
    "[3] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[J]. arXiv preprint arXiv:2004.10934, 2020.",
    "[4] Terven J, Córdova-Esparza D M, Romero-González J A. A Comprehensive Review of YOLO Architectures in Computer Vision: From YOLOv1 to YOLOv8 and YOLO-NAS[J]. Machine Learning and Knowledge Extraction, 2023, 5(4): 1680-1716.",
    "[5] He K, Zhang X, Ren S, Sun J. Deep Residual Learning for Image Recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 770-778.",
    "[6] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]//Advances in Neural Information Processing Systems. Vancouver: Curran Associates, 2019: 8024-8035.",
    "[7] Schroff F, Kalenichenko D, Philbin J. FaceNet: A Unified Embedding for Face Recognition and Clustering[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Boston: IEEE, 2015: 815-823.",
    "[8] Deng J, Guo J, Xue N, Zafeiriou S. ArcFace: Additive Angular Margin Loss for Deep Face Recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Long Beach: IEEE, 2019: 4690-4699.",
    "[9] Ultralytics. YOLOv8[EB/OL]. (2023)[2026-03-28]. https://github.com/ultralytics/ultralytics.",
    "[10] Tran D, Bourdev L, Fergus R, Torresani L, Paluri M. Learning Spatiotemporal Features with 3D Convolutional Networks[C]//Proceedings of the IEEE International Conference on Computer Vision. Santiago: IEEE, 2015: 4489-4497.",
    "[11] Wang L, Xiong Y, Wang Z, et al. Temporal Segment Networks: Towards Good Practices for Deep Action Recognition[C]//Proceedings of the European Conference on Computer Vision. Amsterdam: Springer, 2016: 20-36.",
    "[12] 祝智庭, 胡姣. 教育数字化转型的理论框架[J]. 中国教育学刊, 2022(4): 41-49.",
    "[13] 黄荣怀. 智慧教育的内涵及其特征[J]. 现代远程教育研究, 2014, 26(1): 3-11.",
    "[14] 教育部. 教育信息化2.0行动计划[Z]. 2018-04-13.",
    "[15] Zhang K, Zhang Z, Li Z, Qiao Y. Joint Face Detection and Alignment Using Multitask Cascaded Convolutional Networks[J]. IEEE Signal Processing Letters, 2016, 23(10): 1499-1503.",
    "[16] Kingma D P, Ba J. Adam: A Method for Stochastic Optimization[C]//Proceedings of the International Conference on Learning Representations. San Diego, 2015.",
    "[17] Bradski G, Kaehler A. Learning OpenCV: Computer Vision with the OpenCV Library[M]. Sebastopol: O'Reilly Media, 2008.",
    "[18] Loshchilov I, Hutter F. Decoupled Weight Decay Regularization[C]//Proceedings of the International Conference on Learning Representations. New Orleans: OpenReview.net, 2019.",
    "[19] 飞书开放平台. 自定义机器人使用指南[EB/OL]. [2026-03-29]. https://open.feishu.cn/document/client-docs/bot-v3/add-custom-bot.",
    "[20] 闫兴亚, 匡娅茜, 白光睿, 李月. 基于深度学习的学生课堂行为识别方法[J]. 计算机工程与应用, 2023, 49(7): 251-258.",
    "[21] 郭俊奇, 吕嘉昊, 王汝涵, 熊青云, 张世峰, 胡康颖. 深度学习模型驱动的师生课堂行为识别[J]. 北京师范大学学报(自然科学版), 2021, 57(6): 905-912.",
    "[22] 曾钰琦, 刘博, 钟柏昌, 钟瑾. 智慧教育下基于改进 YOLOv8 的学生课堂行为检测算法[J]. 计算机工程, 2024, 50(9): 344-355.",
    "[23] 全国人民代表大会常务委员会. 中华人民共和国个人信息保护法[EB/OL]. (2021-08-20)[2025-10-31]. https://www.npc.gov.cn/npc/c30834/202108/a8c4e3672c74491a80b53a172bb753fe.shtml.",
    "[24] DeepSeek. API 文档：对话补全（OpenAI 兼容）[EB/OL]. [2026-03-29]. https://api-docs.deepseek.com/zh-cn/.",
]

for ref in references:
    p = add_formatted_paragraph(doc, ref, '宋体', 'Times New Roman', '小四', False,
                               WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)

doc.add_page_break()

# ==================== 8. 致谢 ====================
print("处理致谢...")

p = add_formatted_paragraph(doc, "致 谢", '黑体', 'Times New Roman', '四号', False,
                           WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=20)
doc.add_paragraph()  # 空行

acknowledge = "感谢指导老师【请填写】在选题论证、技术路线与章节安排上的具体建议，使本人得以把课堂所学落到可运行的系统与文字里。感谢信息工程学院各门课程教师的讲授，感谢同窗在调试与排版上的协助，也感谢家人在撰写阶段的包容。学力所限，错漏之处恳请师长指正。"
add_formatted_paragraph(doc, acknowledge, '宋体', 'Times New Roman', '小四', False,
                       WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20)

# ==================== 保存文件 ====================
output_path = r'd:\SY_BYSJ\smart_classroom2\docs\毕业论文_智慧教室系统_已排版.docx'
doc.save(output_path)
print(f"\n排版完成！文件已保存至: {output_path}")
print("请在 Word 中打开检查并刷新目录（引用 → 更新目录）")
